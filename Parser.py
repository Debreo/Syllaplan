import sys
from docx import Document

def main():

    file = "Syllabus.docx"
    document = Document(file)
    tables = document.tables
    print "Extracting Text from Syllabus"
    print ""
    print "*"*10
    for row in tables[2].rows:
        print row.cells[0].paragraphs[0].text + row.cells[1].paragraphs[0].text







if __name__ == "__main__":
    main()

